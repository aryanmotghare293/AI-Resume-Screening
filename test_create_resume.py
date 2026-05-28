"""Generate a sample DOCX resume for testing the app."""
from docx import Document
from docx.shared import Pt, Inches

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

doc.add_heading('Rahul Sharma', level=0)
doc.add_paragraph('Email: rahul.sharma@email.com | Phone: +91-9876543210 | LinkedIn: linkedin.com/in/rahulsharma')

doc.add_heading('Professional Summary', level=1)
doc.add_paragraph(
    'Results-driven Data Scientist with 2+ years of experience in machine learning, '
    'deep learning, and data analytics. Proficient in Python, SQL, and cloud technologies. '
    'Passionate about leveraging AI to solve real-world problems.'
)

doc.add_heading('Education', level=1)
doc.add_paragraph('B.Tech in Computer Science — XYZ University (2020-2024) — CGPA: 8.5/10')

doc.add_heading('Skills', level=1)
doc.add_paragraph(
    'Python, SQL, Machine Learning, Deep Learning, TensorFlow, Scikit-learn, '
    'Pandas, NumPy, Power BI, Git, Docker, REST API, Flask, Streamlit, '
    'Data Visualization, Statistical Modeling, NLP'
)

doc.add_heading('Experience', level=1)
doc.add_paragraph('Data Science Intern — ABC Corp (June 2023 – Dec 2023)')
doc.add_paragraph('• Built churn prediction model')
doc.add_paragraph('• Analyzed customer data')
doc.add_paragraph('• Created dashboards in Power BI')
doc.add_paragraph('ML Engineer Intern — DEF Technologies (Jan 2024 – May 2024)')
doc.add_paragraph('• Worked on NLP project')
doc.add_paragraph('• Developed recommendation system')

doc.add_heading('Projects', level=1)
doc.add_paragraph('1. Customer Churn Prediction — Used logistic regression and XGBoost to predict churn')
doc.add_paragraph('2. Sentiment Analysis Tool — Built an NLP pipeline using BERT for tweet sentiment classification')
doc.add_paragraph('3. Sales Dashboard — Created interactive Power BI dashboard for sales analytics')

doc.add_heading('Certifications', level=1)
doc.add_paragraph('• Google Data Analytics Certificate')
doc.add_paragraph('• Deep Learning Specialization (Coursera)')

doc.save('sample_resume.docx')
print('✅ sample_resume.docx created successfully!')
